import os

import PyPDF2, docx
from docx.shared import Pt

# PDF Paranoia
# Using the os.walk() function from Chapter 9, write a script that will 
# go through every PDF in a folder (and its subfolders) and encrypt the 
# PDFs using a password provided on the command line. Save each encrypted 
# PDF with an _encrypted.pdf suffix added to the original filename. 
# Before deleting the original file, have the program attempt to read 
# and decrypt the file to ensure that it was encrypted correctly.

def PDF_paranoia(path):
    os.chdir(path)
    for root, subfolder, filenames in os.walk(path):
        if 'encrypted' in subfolder:    #exclude output folder
            subfolder.remove('encrypted')                
        for file in filenames:
            if os.path.splitext(file)[1] == '.pdf':  #work with PDF file only
                with open(file, 'rb') as pdf:
                    input_pdf = PyPDF2.PdfFileReader(pdf, strict=False)
                    output_pdf = PyPDF2.PdfFileWriter()
                    output_pdf.appendPagesFromReader(input_pdf)
                    output_pdf.encrypt('abc1234') #encrypt password
                
                #write new PDF file
                with open(os.path.join('encrypted',os.path.splitext(file)[0]+'_encrypted.pdf'),'wb') as pdf_e:
                    print(f'encrypting {os.path.splitext(file)[0]}')
                    output_pdf.write(pdf_e)  

                #Check if PDF existed
                checker = PyPDF2.PdfFileReader(os.path.join('encrypted',os.path.splitext(file)[0]+'_encrypted.pdf'))
                if checker.isEncrypted == True:   #check if PDF is encrypt
                    print(f'confirm {os.path.splitext(file)[0]} encrypted successfully!')
                else:
                    print(f'{os.path.splitext(file)[0]} is still not encrypted!')

                #delete original PDF
                print(f'deleting file {file}')
                os.remove(file)

# Custom Invitations as Word Documents
def custom_Doc_Invitations(path):
    os.chdir(path)
    doc = docx.Document('template invitation.docx')
    with open('guest.txt' ,'r') as f:
        for name in f.readlines():
            name = name.strip() #remove new line
            doc.add_paragraph('It would be a pleasure to have the company of ', style='Style1')
            doc.add_paragraph(name, style = 'Style2')
            doc.add_paragraph('at Central on the Evening of 11st May 2019 at 7 o’clock', style = 'Style3')
            doc.add_page_break()
        doc.save('test.docx') 
        
if __name__ == "__main__":
    path = r'C:\Users\Fung\Downloads\testPDF'
    # PDF_paranoia(path)
    path = r'C:\Users\Fung\Downloads\testDoc'
    # custom_Doc_Invitations(path)
